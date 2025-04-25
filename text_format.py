from docx.shared import Pt

def add_formatted_text_to_document(document, text, rule_number=None):
    lines = text.split('\n')
    for style in document.styles:
        if hasattr(style, 'paragraph_format'):
            style.paragraph_format.line_spacing = 1.0
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(3)

    for i, line in enumerate(lines):
        if line.startswith('**') and i == 0:
            title_text = line.strip('**')
            p = document.add_paragraph()
            if rule_number is not None:
                run = p.add_run(f"{rule_number}. {title_text}")
            else:
                run = p.add_run(title_text)
            run.bold = True
        elif line.startswith('Description:'):
            document.add_paragraph(line)
        elif line.startswith('- '):
            document.add_paragraph(line)
        elif line.startswith('  • '):
            p = document.add_paragraph(line.strip('  • '), style='List Bullet')
        else:
            if line.strip():
                document.add_paragraph(line)

        if 'p' in locals():
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)

    if rule_number is not None:
        p = document.add_paragraph('_' * 105)
        p.paragraph_format.space_after = Pt(9)
        p.alignment = 1  